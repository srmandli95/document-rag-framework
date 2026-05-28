from pathlib import Path

from docx import Document as DocxDocument


def extract_docx_text(file_path: str) -> str:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    document = DocxDocument(str(path))
    extracted_parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            extracted_parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if row_values:
                extracted_parts.append(" | ".join(row_values))

    return "\n".join(extracted_parts)